"""
PDF, Excel, Word ve metin rapor dışa aktarımı — ay bazlı mimari.
"""

from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

import pandas as pd
from matplotlib import get_data_path
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from analyzer import MonthlyAnalysisReport, format_tl
from charts import build_all_charts, chart_display_info
from models import AnalizVeriSeti
from period_utils import analiz_ayi_label

_FONT_REGULAR = "AppSans"
_FONT_BOLD = "AppSans-Bold"
_FONTS_READY = False


def report_heading_base(report: MonthlyAnalysisReport) -> str:
    return f"Finansal Analiz — {analiz_ayi_label(report.analiz_ayi)} (Aylık)"


def _find_turkish_font_files() -> tuple[Path, Path]:
    mpl_fonts = Path(get_data_path()) / "fonts" / "ttf"
    regular, bold = mpl_fonts / "DejaVuSans.ttf", mpl_fonts / "DejaVuSans-Bold.ttf"
    if regular.is_file() and bold.is_file():
        return regular, bold
    if sys.platform == "win32":
        win = Path(r"C:\Windows\Fonts")
        return win / "arial.ttf", win / "arialbd.ttf"
    raise FileNotFoundError("PDF font bulunamadı.")


def _ensure_pdf_fonts() -> None:
    global _FONTS_READY
    if _FONTS_READY:
        return
    regular, bold = _find_turkish_font_files()
    pdfmetrics.registerFont(TTFont(_FONT_REGULAR, str(regular)))
    pdfmetrics.registerFont(TTFont(_FONT_BOLD, str(bold)))
    _FONTS_READY = True


def _pdf_table(data: list[list[str]], col_widths: list[float] | None = None) -> Table:
    t = Table(data, colWidths=col_widths)
    t.setStyle(
        TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), _FONT_BOLD),
            ("FONTNAME", (0, 1), (-1, -1), _FONT_REGULAR),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eaed")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])
    )
    return t


def export_pdf(report: MonthlyAnalysisReport, path: Path) -> None:
    _ensure_pdf_fonts()
    title_style = ParagraphStyle("t", fontName=_FONT_BOLD, fontSize=14, spaceAfter=10)
    body_style = ParagraphStyle("b", fontName=_FONT_REGULAR, fontSize=10, spaceAfter=6)
    story: list = [
        Paragraph(escape(report_heading_base(report)), title_style),
        Paragraph(escape(report.summary_text), body_style),
        Spacer(1, 0.3 * cm),
    ]

    cfo_rows = [["Gösterge", "Değer", "Durum"]]
    for g in report.cfo_uyarilari.gostergeler:
        cfo_rows.append([g.ad, g.deger, g.durum])
    story.append(_pdf_table(cfo_rows, [6 * cm, 4 * cm, 3 * cm]))
    story.append(Spacer(1, 0.3 * cm))

    nakit = report.nakit_akis
    nakit_rows = [
        ["Kalem", "Tutar"],
        ["Dönem Başı", format_tl(nakit.donem_basi_mevcut)],
        ["Girişler", format_tl(nakit.donem_ici_girisler)],
        ["Çıkışlar", format_tl(nakit.donem_ici_cikislar)],
        ["İç Transfer Giriş (102)", format_tl(nakit.ic_transfer_giris)],
        ["İç Transfer Çıkış (102)", format_tl(nakit.ic_transfer_cikis)],
        ["Operasyonel Giriş (transfer hariç)", format_tl(nakit.giris_transfer_haric)],
        ["Operasyonel Çıkış (transfer hariç)", format_tl(nakit.cikis_transfer_haric)],
        ["Dönem Sonu", format_tl(nakit.donem_sonu_net_nakit)],
    ]
    story.append(_pdf_table(nakit_rows, [6 * cm, 4 * cm]))
    story.append(Spacer(1, 0.3 * cm))

    pl = report.aylik_pl
    pl_rows = [
        ["Kalem", "Tutar"],
        ["Brüt Kâr", format_tl(pl.brut_kar)],
        ["Operasyonel Gider", format_tl(pl.toplam_operasyonel_gider)],
        ["Harici Maaş", format_tl(pl.toplam_harici_maas)],
        ["Net Kâr/Zarar", format_tl(pl.aylik_net_kar)],
    ]
    story.append(_pdf_table(pl_rows, [6 * cm, 4 * cm]))
    story.append(Spacer(1, 0.3 * cm))

    if report.ay_karsilastirma:
        k = report.ay_karsilastirma
        onceki = analiz_ayi_label(k.onceki_ayi)
        kiyas_rows = [
            ["Metrik", "Bu Ay", onceki, "Fark"],
            ["Net Kâr/Zarar", format_tl(k.net_kar), format_tl(k.net_kar_onceki), format_tl(k.net_kar_fark)],
            ["Brüt Kâr", format_tl(k.brut_kar), format_tl(k.brut_kar_onceki), format_tl(k.brut_kar_fark)],
            ["Operasyonel Gider", format_tl(k.gider), format_tl(k.gider_onceki), format_tl(k.gider_fark)],
            ["Dönem Sonu Nakit", format_tl(k.nakit_sonu), format_tl(k.nakit_sonu_onceki), format_tl(k.nakit_sonu_fark)],
        ]
        story.append(Paragraph(escape("Önceki Ay Kıyası"), title_style))
        story.append(_pdf_table(kiyas_rows, [4 * cm, 3 * cm, 3 * cm, 3 * cm]))
        story.append(Spacer(1, 0.3 * cm))

    if report.plan_mutabakat:
        pm = report.plan_mutabakat
        mut_rows = [["Kaynak", "Plan", "Yaşlandırma", "Fark"]]
        if pm.tahsilat_plan_acik > 0 or pm.tahsil_edilemeyen > 0:
            mut_rows.append([
                "Tahsilat (120)",
                format_tl(pm.tahsilat_plan_acik),
                format_tl(pm.tahsil_edilemeyen),
                format_tl(pm.tahsilat_fark),
            ])
        if pm.odeme_plan_acik > 0 or pm.odenmeyen > 0:
            mut_rows.append([
                "Ödeme (320)",
                format_tl(pm.odeme_plan_acik),
                format_tl(pm.odenmeyen),
                format_tl(pm.odeme_fark),
            ])
        story.append(Paragraph(escape("Plan Mutabakatı"), title_style))
        story.append(_pdf_table(mut_rows, [4 * cm, 3 * cm, 3 * cm, 3 * cm]))
        story.append(Spacer(1, 0.3 * cm))

    if report.tahsilat_plan or report.tediye_plan:
        plan_rows = [["Plan", "Açık", "Vadesi Geçen", "Vadesi Gelmeyen"]]
        if report.tahsilat_plan and report.tahsilat_plan.hesap_sayisi:
            tp = report.tahsilat_plan
            plan_rows.append([
                "120 Tahsilat",
                format_tl(tp.toplam_acik),
                format_tl(tp.vadesi_gecen),
                format_tl(tp.vadesi_gelmeyen),
            ])
        if report.tediye_plan and report.tediye_plan.hesap_sayisi:
            op = report.tediye_plan
            plan_rows.append([
                "320 Ödeme",
                format_tl(op.toplam_acik),
                format_tl(op.vadesi_gecen),
                format_tl(op.vadesi_gelmeyen),
            ])
        if report.vade_net is not None:
            plan_rows.append(["Vade Neti", format_tl(report.vade_net), "", ""])
        story.append(_pdf_table(plan_rows, [4 * cm, 3.5 * cm, 3.5 * cm, 3.5 * cm]))
        story.append(Spacer(1, 0.3 * cm))

    for rec in report.recommendations[:10]:
        story.append(Paragraph(escape(f"[{rec.oncelik}] {rec.baslik}: {rec.aciklama}"), body_style))

    chart_map = build_all_charts(report)
    if chart_map:
        story.append(Spacer(1, 0.4 * cm))
        story.append(Paragraph(escape("Grafikler"), title_style))
        for key, fig in chart_map.items():
            title, subtitle = chart_display_info(key)
            caption = title + (f" — {subtitle}" if subtitle else "")
            buf = BytesIO()
            fig.savefig(buf, format="png", dpi=120, facecolor=fig.get_facecolor(), bbox_inches="tight")
            buf.seek(0)
            story.append(Paragraph(escape(caption), body_style))
            story.append(Image(buf, width=16 * cm, height=9 * cm))
            story.append(Spacer(1, 0.2 * cm))

    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=1.5 * cm, rightMargin=1.5 * cm)
    doc.build(story)


def _write_sheet(writer: pd.ExcelWriter, name: str, rows: list[list], headers: list[str]) -> None:
    df = pd.DataFrame(rows, columns=headers)
    df.to_excel(writer, sheet_name=name[:31], index=False)


def export_excel(report: MonthlyAnalysisReport, veri: AnalizVeriSeti, path: Path) -> None:
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        cfo = report.cfo_uyarilari
        _write_sheet(
            writer, "CFO",
            [[g.ad, g.deger, g.durum, g.aciklama] for g in cfo.gostergeler],
            ["Gösterge", "Değer", "Durum", "Açıklama"],
        )
        nakit = report.nakit_akis
        _write_sheet(
            writer, "Nakit Akis",
            [
                ["Dönem Başı", nakit.donem_basi_mevcut],
                ["Girişler", nakit.donem_ici_girisler],
                ["Çıkışlar", nakit.donem_ici_cikislar],
                ["İç Transfer Giriş (102)", nakit.ic_transfer_giris],
                ["İç Transfer Çıkış (102)", nakit.ic_transfer_cikis],
                ["Operasyonel Giriş (transfer hariç)", nakit.giris_transfer_haric],
                ["Operasyonel Çıkış (transfer hariç)", nakit.cikis_transfer_haric],
                ["Dönem Sonu", nakit.donem_sonu_net_nakit],
            ],
            ["Kalem", "Tutar"],
        )
        pl = report.aylik_pl
        _write_sheet(
            writer, "Aylik PL",
            [
                ["Brüt Kâr", pl.brut_kar],
                ["Operasyonel Gider", pl.toplam_operasyonel_gider],
                ["Harici Maaş", pl.toplam_harici_maas],
                ["Net", pl.aylik_net_kar],
            ],
            ["Kalem", "Tutar"],
        )
        if report.ay_karsilastirma:
            k = report.ay_karsilastirma
            onceki = analiz_ayi_label(k.onceki_ayi)
            _write_sheet(
                writer, "Onceki Ay Kiyasi",
                [
                    ["Net Kâr/Zarar", k.net_kar, k.net_kar_onceki, k.net_kar_fark],
                    ["Brüt Kâr", k.brut_kar, k.brut_kar_onceki, k.brut_kar_fark],
                    ["Operasyonel Gider", k.gider, k.gider_onceki, k.gider_fark],
                    ["Dönem Sonu Nakit", k.nakit_sonu, k.nakit_sonu_onceki, k.nakit_sonu_fark],
                ],
                ["Metrik", "Bu Ay", onceki, "Fark"],
            )
        if report.plan_mutabakat:
            pm = report.plan_mutabakat
            mut_rows: list[list] = []
            if pm.tahsilat_plan_acik > 0 or pm.tahsil_edilemeyen > 0:
                mut_rows.append([
                    "Tahsilat (120)",
                    pm.tahsilat_plan_acik,
                    pm.tahsil_edilemeyen,
                    pm.tahsilat_fark,
                ])
            if pm.odeme_plan_acik > 0 or pm.odenmeyen > 0:
                mut_rows.append([
                    "Ödeme (320)",
                    pm.odeme_plan_acik,
                    pm.odenmeyen,
                    pm.odeme_fark,
                ])
            if mut_rows:
                _write_sheet(
                    writer, "Plan Mutabakat",
                    mut_rows,
                    ["Kaynak", "Plan", "Yaşlandırma", "Fark"],
                )
        if not veri.muavin_df.empty:
            veri.muavin_df.to_excel(writer, sheet_name="Muavin", index=False)
        if not veri.banka_df.empty:
            veri.banka_df.to_excel(writer, sheet_name="Banka", index=False)
        if not veri.alis_fatura_df.empty:
            veri.alis_fatura_df.to_excel(writer, sheet_name="Alis Fatura", index=False)
        if not veri.satis_fatura_df.empty:
            veri.satis_fatura_df.to_excel(writer, sheet_name="Satis Fatura", index=False)
        if not veri.tahsilat_plan_df.empty:
            veri.tahsilat_plan_df.to_excel(writer, sheet_name="Tahsilat Plan", index=False)
        if not veri.tediye_plan_df.empty:
            veri.tediye_plan_df.to_excel(writer, sheet_name="Tediye Plan", index=False)
        if report.tahsilat_plan and report.tahsilat_plan.top_hesaplar:
            _write_sheet(
                writer, "Tahsilat Ozet",
                [
                    [h.hesap_kodu, h.hesap_adi, h.acik_bakiye, h.vadesi_gecen, h.vadesi_gelmeyen]
                    for h in report.tahsilat_plan.top_hesaplar
                ],
                ["Cari", "Ad", "Açık", "Vadesi Geçen", "Vadesi Gelmeyen"],
            )
        if report.tediye_plan and report.tediye_plan.top_hesaplar:
            _write_sheet(
                writer, "Tediye Ozet",
                [
                    [h.hesap_kodu, h.hesap_adi, h.acik_bakiye, h.vadesi_gecen, h.vadesi_gelmeyen]
                    for h in report.tediye_plan.top_hesaplar
                ],
                ["Cari", "Ad", "Açık", "Vadesi Geçen", "Vadesi Gelmeyen"],
            )
        _write_sheet(
            writer, "Yaslandirma",
            [
                [s.cari_kodu, s.cari_adi, s.fatura_tutari, s.eslesen_banka, s.kalan, s.bucket]
                for s in report.yaslandirma.tahsil_edilemeyen + report.yaslandirma.odenmeyen
            ],
            ["Karşı Hesap", "Cari Adı", "Fatura", "Banka", "Kalan", "Bucket"],
        )
        _write_sheet(
            writer, "Eslesme",
            [
                [g.tarih, g.banka_adi, g.cari_kodu, g.kategori, g.aciklama, g.tutar]
                for g in report.eslesme.aciklanamayan_giderler
            ],
            ["Tarih", "Banka", "Karşı Hesap", "Kategori", "Açıklama", "Tutar"],
        )
        _write_sheet(
            writer, "Tavsiyeler",
            [[r.kategori, r.oncelik, r.baslik, r.aciklama] for r in report.recommendations],
            ["Kategori", "Öncelik", "Başlık", "Açıklama"],
        )
        for sheet in writer.book.worksheets:
            for col in sheet.columns:
                sheet.column_dimensions[get_column_letter(col[0].column)].width = 18
            sheet["A1"].font = Font(bold=True)


def export_word(report: MonthlyAnalysisReport, path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(report_heading_base(report), 0)
    doc.add_paragraph(report.summary_text)
    doc.add_heading("CFO Erken Uyarı", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Gösterge", "Değer", "Durum"
    for g in report.cfo_uyarilari.gostergeler:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = g.ad, g.deger, g.durum
    doc.add_heading("Nakit Akış", level=1)
    doc.add_paragraph(report.nakit_akis.yorum)
    doc.add_heading("Aylık P&L", level=1)
    pl = report.aylik_pl
    for line in [
        f"Brüt kâr: {format_tl(pl.brut_kar)}",
        f"Gider: {format_tl(pl.toplam_operasyonel_gider)}",
        f"Net: {format_tl(pl.aylik_net_kar)}",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    if report.tahsilat_plan or report.tediye_plan:
        doc.add_heading("Tahsilat ve Ödeme Planı", level=1)
        for sec in report.summary_sections:
            if sec.baslik == "Tahsilat ve Ödeme Planı":
                for line in sec.satirlar:
                    doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("Tavsiyeler", level=1)
    for rec in report.recommendations[:15]:
        doc.add_paragraph(f"[{rec.oncelik}] {rec.baslik}: {rec.aciklama}")
    doc.save(path)


def export_text(report: MonthlyAnalysisReport, path: Path) -> None:
    lines = [
        report_heading_base(report),
        "=" * len(report_heading_base(report)),
        "",
        report.summary_text,
        "",
        "CFO ERKEN UYARI",
        report.cfo_uyarilari.ozet_metin,
        "",
        "NAKİT AKIŞ",
        report.nakit_akis.yorum,
        "",
        "AYLIK P&L",
        f"Brüt kâr: {format_tl(report.aylik_pl.brut_kar)}",
        f"Net: {format_tl(report.aylik_pl.aylik_net_kar)}",
        "",
    ]
    for sec in report.summary_sections:
        if sec.baslik == "Tahsilat ve Ödeme Planı":
            lines.append("TAHSİLAT VE ÖDEME PLANI")
            lines.extend(sec.satirlar)
            lines.append("")
            break
    lines.extend([
        "TAVSİYELER",
    ])
    for rec in report.recommendations:
        lines.append(f"[{rec.oncelik}] {rec.baslik}: {rec.aciklama}")
    path.write_text("\n".join(lines), encoding="utf-8")
